"""
File Processing Service for In-Memory Document Handling
Supports: PDF, DOCX, DOC, TXT
"""

import io
import pdfplumber
from docx import Document
from werkzeug.utils import secure_filename
import mimetypes

class FileProcessingService:
    """Service to extract text from uploaded files in-memory"""
    
    ALLOWED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}
    ALLOWED_MIMETYPES = {
        'text/plain',
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    }
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @classmethod
    def validate_file(cls, file, filename):
        """
        Validate file type and size
        
        Args:
            file: FileStorage object
            filename: Original filename
            
        Returns:
            dict: {'valid': bool, 'error': str or None, 'sanitized_name': str}
        """
        if not file or not filename:
            return {'valid': False, 'error': 'No file provided', 'sanitized_name': None}
        
        # Sanitize filename
        sanitized_name = secure_filename(filename)
        if not sanitized_name:
            return {'valid': False, 'error': 'Invalid filename', 'sanitized_name': None}
        
        # Check file extension
        file_ext = '.' + sanitized_name.rsplit('.', 1)[1].lower() if '.' in sanitized_name else ''
        if file_ext not in cls.ALLOWED_EXTENSIONS:
            return {
                'valid': False, 
                'error': f'File type not allowed. Allowed: {", ".join(cls.ALLOWED_EXTENSIONS)}',
                'sanitized_name': sanitized_name
            }
        
        # Check MIME type
        mimetype = file.content_type or mimetypes.guess_type(filename)[0]
        if mimetype not in cls.ALLOWED_MIMETYPES:
            return {
                'valid': False,
                'error': f'Invalid MIME type: {mimetype}',
                'sanitized_name': sanitized_name
            }
        
        # Check file size (read and seek back)
        file.seek(0, io.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > cls.MAX_FILE_SIZE:
            return {
                'valid': False,
                'error': f'File too large. Max size: {cls.MAX_FILE_SIZE / 1024 / 1024}MB',
                'sanitized_name': sanitized_name
            }
        
        if file_size == 0:
            return {
                'valid': False,
                'error': 'File is empty',
                'sanitized_name': sanitized_name
            }
        
        return {'valid': True, 'error': None, 'sanitized_name': sanitized_name}
    
    @classmethod
    def extract_text_from_pdf(cls, file_bytes):
        """
        Extract text from PDF file in-memory
        
        Args:
            file_bytes: BytesIO object containing PDF data
            
        Returns:
            str: Extracted text
        """
        try:
            text_content = []
            with pdfplumber.open(file_bytes) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---\n{page_text}")
            
            extracted_text = "\n\n".join(text_content)
            
            if not extracted_text.strip():
                return "⚠️ PDF file appears to be empty or contains only images."
            
            return extracted_text
            
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
    
    @classmethod
    def extract_text_from_docx(cls, file_bytes):
        """
        Extract text from DOCX/DOC file in-memory
        
        Args:
            file_bytes: BytesIO object containing DOCX data
            
        Returns:
            str: Extracted text
        """
        try:
            doc = Document(file_bytes)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    table_texts.append("\n".join(table_data))
            
            # Combine all text
            all_text = []
            if paragraphs:
                all_text.append("\n\n".join(paragraphs))
            if table_texts:
                all_text.append("\n\n--- Tables ---\n\n" + "\n\n".join(table_texts))
            
            extracted_text = "\n\n".join(all_text)
            
            if not extracted_text.strip():
                return "⚠️ Document appears to be empty."
            
            return extracted_text
            
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    @classmethod
    def extract_text_from_txt(cls, file_bytes):
        """
        Extract text from TXT file in-memory
        
        Args:
            file_bytes: BytesIO object containing text data
            
        Returns:
            str: Extracted text
        """
        try:
            # Try UTF-8 first, then fall back to other encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    file_bytes.seek(0)
                    text = file_bytes.read().decode(encoding)
                    
                    if not text.strip():
                        return "⚠️ Text file appears to be empty."
                    
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with supported encodings")
            
        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")
    
    @classmethod
    def process_file(cls, file, filename):
        """
        Main method to process uploaded file and extract text
        
        Args:
            file: FileStorage object from Flask request
            filename: Original filename
            
        Returns:
            dict: {
                'success': bool,
                'extracted_text': str or None,
                'filename': str,
                'file_type': str,
                'file_size': int,
                'error': str or None
            }
        """
        # Validate file
        validation = cls.validate_file(file, filename)
        if not validation['valid']:
            return {
                'success': False,
                'extracted_text': None,
                'filename': filename,
                'file_type': None,
                'file_size': 0,
                'error': validation['error']
            }
        
        sanitized_name = validation['sanitized_name']
        file_ext = '.' + sanitized_name.rsplit('.', 1)[1].lower()
        
        try:
            # Read file into memory
            file_bytes = io.BytesIO(file.read())
            file_size = len(file_bytes.getvalue())
            
            # Extract text based on file type
            if file_ext == '.pdf':
                extracted_text = cls.extract_text_from_pdf(file_bytes)
                file_type = 'pdf'
            elif file_ext in ['.docx', '.doc']:
                extracted_text = cls.extract_text_from_docx(file_bytes)
                file_type = 'docx'
            elif file_ext == '.txt':
                extracted_text = cls.extract_text_from_txt(file_bytes)
                file_type = 'txt'
            else:
                raise Exception(f"Unsupported file type: {file_ext}")
            
            # Clear the BytesIO object
            file_bytes.close()
            
            return {
                'success': True,
                'extracted_text': extracted_text,
                'filename': sanitized_name,
                'file_type': file_type,
                'file_size': file_size,
                'error': None
            }
            
        except Exception as e:
            return {
                'success': False,
                'extracted_text': None,
                'filename': sanitized_name,
                'file_type': None,
                'file_size': 0,
                'error': str(e)
            }

# Export singleton instance
file_processor = FileProcessingService()